from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(11)

# ── 标题 ──
title = doc.add_heading('今日工作总结与思考', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(f'2026年7月3日 · 周五')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(128, 128, 128)

doc.add_paragraph()

# ══════ 第一部分 ══════
doc.add_heading('一、今日工作成果', level=1)

sections = [
    ("JWT 登录系统", 
     "今天完成了系统最重要的安全基础设施——JWT 认证登录。我设计了完整的认证链路："
     "用户通过独立的全屏登录页输入账号密码 → 后端 bcrypt 验证 → 签发 24 小时 JWT token → "
     "前端 localStorage 存储 → axios 拦截器自动附加 Bearer token → 路由守卫检查未登录跳转。"
     "登录页面本身投入了大量精力，用 Canvas 粒子背景、CSS 轨道环动画、玻璃态卡片、"
     "3D 视差效果打造了一个商业级品质的登录体验。"),
    
    ("权限管理系统",
     "实现了管理员/普通用户双角色权限体系。普通用户登录后只能看到导航框架，"
     "内容区被半透明遮罩覆盖，需要向管理员申请权限。管理员在顶栏看到待审批红色角标，"
     "进入权限管理页面可以同意或拒绝。整个审批流的 UI 和 API 都是今天从零搭建的。"),
    
    ("仪表盘模块化重构",
     "发现仪表盘的库存饼图按 material_type 分组，导致大部分零件都被归为'外购件'，"
     "其他模块几乎不可见。我重写了统计逻辑，通过 BOM 线反查物料所属模块，"
     "将饼图改为横向条形图，让外购件、电气、视觉、量具工具等每个模块平等展示。"),
    
    ("P0 安全加固",
     "完成了三项关键安全措施：JWT Secret 改用环境变量（未设则随机生成）、"
     "smtp_config.json 加入 .gitignore 防止授权码泄露、requirements.txt 补全 "
     "python-jose / passlib / bcrypt 依赖并锁定 bcrypt 4.0.1 版本。"),
    
    ("邮件通知与定时 MRP",
     "完善了邮件通知功能：notifier.py 新增 JSON 持久化，SMTP 配置重启不丢失。"
     "前端系统工具 popover 增加邮件配置面板。定时 MRP 默认启用每天 06:00 自动运算。"
     "修复了输入框暗色不可见 bug——折腾了三轮才搞定——最终用原生 input 替换 Element Plus 组件。"),
    
    ("云端系统全功能演练 + 7 问题修复",
     "以管理员身份对 Railway 生产系统执行了完整的全功能演练：登录 → "
     "MRP 运算(69 条计划订单) → 创建销售订单 → 销售转 MPS → 再跑 MRP。"
     "发现了 7 个问题并全部修复：销售订单 API 加 Pydantic 模型、创建 user1 普通用户、"
     "降低种子数据安全库存、MRP 转单空数组提示、默认启用定时器、修复工单编号 bug。"),
    
    ("云端数据迁移",
     "最折腾的一步。Railway 的 import API 在 SQLite 上多次失败——调试了四轮："
     "先是 Pydantic 解析问题，然后是外键约束导致 DELETE 失败，接着不小心清空了 users 表，"
     "最后修正了保留系统表 + 保留 ID 引用的逻辑。最终成功将 178 种物料、177 行 BOM、"
     "173 条库存的完整三工位数据导入到 Railway 云端。"),
    
    ("文档体系搭建",
     "创建了项目 README，重写了开发文档（新增认证权限/定时MRP/邮件章节）、"
     "用户手册（新增登录权限/邮件配置/FAQ）、测试指南（新增 JWT/权限测试用例）。"
     "还制作了一份 HTML 格式的云端操作指南和配套的演示测试数据集。"),
]

for title, body in sections:
    doc.add_heading(title, level=2)
    doc.add_paragraph(body)

# ══════ 第二部分 ══════
doc.add_heading('二、今日技术思考', level=1)

thoughts = [
    ("关于 JWT 认证的设计取舍",
     "选择了 single-admin-seed 模式（数据库启动时自动创建默认账号），而不是开放注册。"
     "这对于企业级 MRP 系统是正确的——用户应该由管理员创建和管理，而非自主注册。"
     "24 小时的 token 过期时间对内部系统来说也合理，不会太频繁打扰用户。"
     "但后续应该加上密码修改功能和 token refresh 机制。"),
    
    ("Element Plus 暗色主题的坑",
     "定时器输入框的数字显示问题花了很长时间排查。Core issue 是 Element Plus Popover "
     "用 Teleport 渲染到 body 下，scoped CSS 的 :deep() 根本命中不了。"
     "最终换了思路——不用 el-input-number，直接用原生 <input type='number'> 加全局 CSS。"
     "这说明在设计暗色主题系统时，需要提前考虑组件库的 Portal/Teleport 渲染机制。"),
    
    ("Railway SQLite 的数据迁移教训",
     "今天最耗时的就是这一块。核心问题是 import API 代码最初是为 MySQL 写的，"
     "在 SQLite 上 DELETE 时因为外键约束顺序不当而失败。而且每次部署要等 1-2 分钟才能验证。"
     "最大的教训：永远要保留 users 表不动，删数据要按外键依赖逆序且分两轮执行。"
     "以后新建 API 应该先在本地 SQLite 测过再部署。"),
    
    ("权限系统的架构思考",
     "当前的 require_approved 依赖注入放在 main.py 路由注册层，"
     "好处是集中管理，坏处是新加路由可能忘记加保护。更优雅的做法是在每个 router 内部"
     "显式声明依赖。另外，普通用户的 '只能看框架不能看数据' 实现方式（前端遮罩 + "
     "后端 403）是合理的，但遮罩组件目前只监听 user.is_approved，"
     "如果管理员中途撤销权限，普通用户需要刷新才能看到遮罩——后续可以加 WebSocket 推送。"),
    
    ("关于 BOM 架构的反思",
     "目前系统采用单 BOM Header + 多父物料的架构，模块通过 parent_item_id 区分。"
     "这种架构的优点是简单（一个项目一个 BOM），但缺点是代码中到处需要特殊处理——"
     "BOM 树、物料树、成本汇总、采购同步都要根据 parent_item_id 回溯模块。"
     "如果重来，我会选择每个模块独立 BOM Header 的设计，虽然表多了但逻辑更清晰。"),
]

for title, body in thoughts:
    doc.add_heading(title, level=2)
    doc.add_paragraph(body)

# ══════ 第三部分 ══════
doc.add_heading('三、明天待办', level=1)

todos = [
    "补全量具工具类 BOM 的第二张表数据（金山文档中约 60 行工具耗材未导入）",
    "添加用户注册功能（管理员创建新用户）",
    "JWT token refresh 机制",
    "MRP 引擎读在途量/在制量，完善净需求计算",
    "单元测试框架搭建（至少覆盖 MRP 引擎和权限系统）",
    "Docker 部署脚本优化",
]

for t in todos:
    p = doc.add_paragraph(t, style='List Bullet')

# ── 结尾 ──
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("— 写于 2026年7月3日 18:00")
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(128, 128, 128)
run.italic = True

# 保存
doc.save(r'C:\Users\20210817\Desktop\今日工作总结与思考.docx')
print("✅ 已保存到桌面")
