"""Generate today's work summary and technical analysis documents"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

today = datetime.now().strftime("%Y年%m月%d日")

# ========== Doc 1: Work Summary ==========
doc1 = Document()
for s in doc1.sections:
    s.top_margin = Cm(2.5); s.bottom_margin = Cm(2); s.left_margin = Cm(2.5); s.right_margin = Cm(2.5)

doc1.add_heading("MRP II 系统开发工作日报", level=0)
doc1.add_paragraph(f"日期：{today}").alignment = WD_ALIGN_PARAGRAPH.RIGHT
doc1.add_paragraph("项目：三工位-15台测试台 MRP II 系统").alignment = WD_ALIGN_PARAGRAPH.RIGHT
doc1.add_paragraph("")

doc1.add_heading("一、前端 UI 全面重构", level=1)
doc1.add_paragraph("今天对前端界面进行了系统级重构，从 Element Plus 默认模板迁移到紧凑专业型的暗色主题方案，参考 VS Code 设计语言。")
for item in [
    "从 el-menu 迁移到基于 Tailwind CSS 的自定义导航，NavGroup/NavItem 组件化",
    "导航分组：基础数据 / 计划与执行 / 分析与监控 / 财务",
    "新增 Dashboard 仪表板、SearchPalette 全局搜索（Ctrl+K）",
    "面包屑导航、渐变Logo、运行状态脉动指示灯",
    "全页面暗色主题：CSS变量设计令牌体系 + 组件级深色覆盖",
    "按钮降饱和度，导出/导入并排紧凑布局",
    "系统工具从底部固定栏改为底部 Popover 弹窗",
    "移除一键初始化按钮，精简功能入口",
    "键盘快捷键：Ctrl+K 搜索、Ctrl+M MRP、Ctrl+D 仪表板、Ctrl+B BOM",
]:
    doc1.add_paragraph(item, style="List Bullet")

doc1.add_heading("二、数据清空与金山文档导入", level=1)
doc1.add_paragraph("从金山文档「三工位-15台测试.xlsx」读取5个BOM表格数据导入系统。")
for item in [
    "成品：三工位测试台（15台套），编码 SG-3ST-15",
    "5个模块：外购件、外加工件、电气、视觉、量具工具",
    "70个零件（品牌/单价/采购链接完整），三级BOM共75行",
    "MPS计划：3批15台（7天/21天/35天）",
    "费用合计：¥93,250（视觉模块占比最高¥48,990）",
]:
    doc1.add_paragraph(item, style="List Bullet")

doc1.add_heading("三、模块联动修复", level=1)
doc1.add_paragraph("导入数据后多个页面显示「暂无数据」。根因：导入脚本将所有零件挂在同一产品BOM头下，系统代码假设模块有独立BOM头。4个API修复：")
for item in [
    "BOM树查看：seen_headers改为seen_pairs，支持同BOM头递归展开",
    "物料管理树：模块无独立BOM头时回退到产品BOM头查找",
    "费用合计：回退逻辑 + 修复缩进bug",
    "采购同步BOM：3处补上bom_header_id筛选",
    "模块level_type统一为「模块」",
]:
    doc1.add_paragraph(item, style="List Bullet")

doc1.add_heading("四、全流程模拟操作", level=1)
doc1.add_paragraph("按操作手册完成10步业务流程：创建客户→销售订单→MPS→MRP→例外→采购单→到货入库→出库→收款→订单确认。")
doc1.add_paragraph("收据：SO-20260702-0001，出货3/10台，收款¥15000/¥50000，状态「进行中」")
doc1.add_paragraph("发现的问题：到货API路径为 /status 后缀、出入库需 warehouse_id 整数、初始库存为0无法出库、收款需 sales_order_id + customer_id 双重参数")

doc1.add_heading("五、系统功能现状", level=1)
modules = [
    ("物料管理", "三折叠浏览、搜索、新增/编辑、批量修改"),
    ("BOM管理", "树形编辑、导入导出、版本管理、ECN变更"),
    ("库存管理", "出入库、呆滞料预警、流水查询"),
    ("MPS主计划", "手动/批量生成、销售订单追踪"),
    ("MRP运算", "LLC层级+净需求计算、批量规则、替代料"),
    ("采购管理", "BOM同步、到货入库、供应商管理"),
    ("销售管理", "客户/订单管理、三状态联动"),
    ("财务管理", "收款/汇总/订单状态同步"),
    ("费用合计", "三级费用展开"),
    ("例外看板", "缺料/逾期/安全库存预警"),
    ("其他", "生产、工艺路线、CRP产能、检验盘点、报表分析"),
]
for name, desc in modules:
    p = doc1.add_paragraph()
    run = p.add_run(f"{name}：")
    run.bold = True
    p.add_run(desc)

doc1.add_paragraph("")
doc1.add_paragraph(f"系统版本 v1.6.0，FastAPI + MySQL + Vue3 + Element Plus + Tailwind CSS")

out1 = r"C:\Users\20210817\Desktop\今日工作总结.docx"
doc1.save(out1)
print(f"OK: {out1}")

# ========== Doc 2: Technical Analysis ==========
doc2 = Document()
for s in doc2.sections:
    s.top_margin = Cm(2.5); s.bottom_margin = Cm(2); s.left_margin = Cm(2.5); s.right_margin = Cm(2.5)

doc2.add_heading("今日思考与技术分析", level=0)
doc2.add_paragraph(f"日期：{today}").alignment = WD_ALIGN_PARAGRAPH.RIGHT
doc2.add_paragraph("")

doc2.add_heading("一、关于系统架构的反思", level=1)

sections = {
    "1.1 BOM头设计的架构选择": [
        "今天最核心的技术发现：BomHeader 设计存在「双重模式」——既能支持每个模块独立BOM头的分散模式，也能支持单产品BOM头下挂所有层次的集中模式。",
        "导入数据时选择集中模式（一个BOM头75条行），而系统代码在4个地方假设了分散模式。这不是bug，而是「隐式架构约束」没有显式表达。",
        "教训：数据模型中多对多关系天然灵活，但业务逻辑中的隐式约束会在新人操作或重构时断裂。应通过数据校验或统一兜底逻辑消除不一致。",
    ],
    "1.2 测试覆盖的盲区": [
        "三轮测试38个API全部200通过，但没发现BOM树只能展2层、物料树0个零件。因为测试只做「浅度校验」（状态码+非空），没做「结构深度校验」。",
        "对树形接口，正确测试应验证：节点总数、各深度层节点数、叶子节点非空、父子关系一致性。成本不高但发现率大幅提高。",
        "另一个盲区：只测后端API，没模拟前端消费链路。前端调用 /materials/tree 而非 /materials 列表接口，这个差异导致漏测。以后应重点测前端实际消费的端点。",
    ],
    "1.3 导入工具的工程化不足": [
        "金山文档5个sheet共600+行数据，导入脚本只提取了70条。因为脚本适配固定列序格式，与金山文档结构不匹配。",
        "理想导入系统应：支持列映射配置（Sheet→模块，列名→字段）、增量导入和幂等性、导入前预校验预览、完整错误回滚。",
        "当前 excel_importer.py 只支持特定列名格式，与金山文档完全不匹配，下轮迭代需解决。",
    ],
    "1.4 前端重构的观察": [
        "Tailwind替代Element Plus布局组件：获得像素级样式控制权。",
        "CSS变量设计令牌：主题切换成本降到几乎为零。",
        "NavGroup/NavItem 组件化：导航从数据驱动变成组件驱动，可读性提升。",
        "SearchPalette Teleport + 键盘导航：专业级应用特征。",
        "注意：router-link 精确匹配可能在某些嵌套路由下失效；旧全局暗色CSS与新Tailwind令牌可能有冲突需清理。",
    ],
}

for title, paras in sections.items():
    doc2.add_heading(title, level=2)
    for p_text in paras:
        doc2.add_paragraph(p_text)

doc2.add_heading("二、技术债务清单", level=1)
debts = [
    ("P0", "supplier_id硬编码", "所有MRP采购单挂同一供应商。修复：物料default_supplier_id + 查询关联"),
    ("P0", "PO编号并发安全", "seq+1非原子操作。修复：UUID短码或数据库行锁"),
    ("P0", "MRP引擎N+1查询", "每条BOM行2次DB查询。修复：一次查询+内存索引"),
    ("P0", "BOM更新事务保护", "先删后建无事务。修复：try/commit/rollback"),
    ("P1", "API路径不友好", "到货PUT /status 路径冗长，参数不直观。建议简化"),
    ("P1", "无Pydantic验证", "裸dict无类型校验。修复：Pydantic Schema"),
    ("P2", "Alembic迁移", "手动ALTER TABLE。需要自动版本控制"),
    ("P2", "批量文件导入", "API逐行导入极慢。需支持文件上传批量"),
]
for severity, title, desc in debts:
    p = doc2.add_paragraph()
    run = p.add_run(f"{severity} - {title}：")
    run.bold = True
    p.add_run(desc)

doc2.add_heading("三、下一步建议", level=1)
suggestions = [
    "优先修复P0级债务（supplier_id/N+1/事务），直接影响正确性和性能",
    "完善Excel导入，支持金山文档格式列映射和批量导入",
    "加强测试：增加结构深度校验层，覆盖前端实际消费端点",
    "引入Pinia状态管理，解决跨组件数据共享",
    "BOM管理增加可视化树形编辑（拖拽构建BOM树）",
]
for s in suggestions:
    doc2.add_paragraph(s, style="List Number")

doc2.add_heading("四、今日收获", level=1)
insights = [
    "「API返回200不代表功能正确」—— 最深一课。38端点全通过但页面全是「暂无数据」。深度校验不是可选项。",
    "「隐式架构约束」比显式bug更危险 —— 假设模块有独立BOM头，但无文档化或强制校验。",
    "端点命名和参数设计直接影响可发现性 —— 花很时间摸索正确路径。好API应直觉可猜。",
    "三层BOM（产品→模块→零件）正确 —— 费用合计直接展开到模块成本（视觉¥48,990，占52%）。",
    "Vue3 + Element Plus + Tailwind CSS：组件生态丰富 + 设计自由度不受限。生产级前端架构好选。",
]
for insight in insights:
    doc2.add_paragraph(insight, style="List Bullet")

doc2.add_paragraph("")
doc2.add_paragraph(f"记录于 {today}，MRP II 系统项目")

out2 = r"C:\Users\20210817\Desktop\今日思考与技术分析.docx"
doc2.save(out2)
print(f"OK: {out2}")
