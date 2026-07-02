from docx import Document
from docx.shared import Pt
from datetime import datetime

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(10)

doc.add_heading('MRP II 系统编码规范文档', level=0)
doc.add_paragraph(f'版本 v1.0 | 制定日期: {datetime.now().strftime("%Y-%m-%d")} | 适用范围: 前端Vue3 + 后端FastAPI')

doc.add_heading('一、项目结构规范', level=1)
doc.add_paragraph(
    'mrp-system/\n'
    '├── backend/\n'
    '│   ├── app/\n'
    '│   │   ├── api/          # API路由层 (每模块一个文件)\n'
    '│   │   ├── models/       # SQLAlchemy数据模型\n'
    '│   │   ├── services/     # 业务逻辑层 (Service类)\n'
    '│   │   ├── core/         # 基础设施 (database.py, config.py)\n'
    '│   │   └── main.py       # FastAPI入口\n'
    '│   └── seed_data.py      # 测试数据\n'
    '└── frontend/\n'
    '    └── src/\n'
    '        ├── views/        # 页面组件 (PascalCase)\n'
    '        ├── api/          # HTTP请求封装\n'
    '        ├── router/       # Vue Router\n'
    '        └── App.vue       # 根组件'
)

sections = [
    ('2.1 API路由层', [
        '文件名: 小写下划线 (purchase.py, sales.py)',
        '路由前缀: /api/{module_name}',
        '函数名: verb_noun (list_orders, create_bom)',
        '标准CRUD: GET /, GET /{id}, POST /, PUT /{id}, DELETE /{id}',
    ]),
    ('2.2 数据模型层', [
        '类名: PascalCase, 表名: 小写下划线',
        '所有表含 id(PK) + created_at',
        '新增列: Column定义 + init_db() ALTER TABLE迁移',
    ]),
    ('2.3 业务服务层', [
        '跨模块逻辑抽取为Service类 (@staticmethod)',
        '已有: SalesOrderService, BomExploder, MrpCalculator',
    ]),
    ('2.4 数据库迁移', [
        '新增列必须在 init_db() migration_stmts 中添加',
        '格式: ALTER TABLE {t} ADD COLUMN {c} {type} DEFAULT {v}',
        '每条独立 try/except, 标注版本号',
    ]),
]

doc.add_heading('二、后端编码规范', level=1)
for title, items in sections:
    doc.add_heading(title, level=2)
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

frontend_sections = [
    ('3.1 页面组件', [
        '命名: PascalCase (PurchaseList.vue, CostSummary.vue)',
        '结构: <template> -> <script setup> -> <style scoped>',
        '工具条: .page-toolbar, display:flex, gap:12px',
        '搜索框: 240-280px, @clear + @keyup.enter',
    ]),
    ('3.2 三级折叠视图', [
        'CSS: .proj-card/.proj-header, .mod-card/.mod-header',
        '箭头: ▲展开/▼收起, 背景#fafbfc, hover #f0f5ff',
        '状态: 独立Set/Map, 不用computed内_open',
        '批量栏: .batch-bar, 背景#f0f9eb, 边框#b3e19d',
    ]),
    ('3.3 API请求', [
        '统一 import api from "@/api"',
        'GET api.get(url,{params}) / POST api.post(url,data)',
        '分页: 全量用 /all, 其他统一page_size=1000',
        '禁止硬编码自定义page_size',
    ]),
    ('3.4 状态管理', [
        'ref(): 简单值 / reactive(): 对象表单',
        '多表格选择: 独立映射, 不互相覆盖',
        '数据源: 预缓存到 _enriched, 避免每次创建新对象',
    ]),
]

doc.add_heading('三、前端编码规范', level=1)
for title, items in frontend_sections:
    doc.add_heading(title, level=2)
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

doc.add_heading('四、API端点命名规范', level=1)
endpoints = [
    'GET /{resource} -> 列表 / GET /{resource}/{id} -> 详情',
    'POST /{resource} -> 创建 / PUT /{resource}/{id} -> 更新',
    'DELETE /{resource}/{id} -> 删除',
    '特殊: /all(全量), /summary(汇总), /batch(批量), /tree/{id}(树)',
]
for ep in endpoints:
    doc.add_paragraph(ep, style='List Bullet')

doc.add_heading('五、错误处理', level=1)
errors = [
    '后端: HTTPException (404/400/500)',
    '前端: try/catch + ElMessage.error()',
    '危险操作: ElMessageBox.confirm()',
]
for e in errors:
    doc.add_paragraph(e, style='List Bullet')

doc.add_heading('六、Git规范', level=1)
git_rules = [
    '版本号: v{major}.{minor}.{patch}',
    'Commit: v1.5.0: 标题 + 详细说明 + 文件统计',
    '推送: git push origin main -> Railway自动部署',
]
for g in git_rules:
    doc.add_paragraph(g, style='List Bullet')

doc.add_heading('七、命名速查表', level=1)
t = doc.add_table(rows=12, cols=3, style='Table Grid')
h = t.rows[0].cells; h[0].text = '类别'; h[1].text = '规范'; h[2].text = '示例'
data = [
    ('Python文件', '小写_下划线', 'purchase.py'),
    ('Python类', 'PascalCase', 'SalesOrderService'),
    ('Python函数', 'verb_noun', 'list_orders'),
    ('SQL表', '小写_下划线', 'purchase_order'),
    ('SQL列', '小写_下划线', 'order_qty'),
    ('Vue组件', 'PascalCase', 'PurchaseList.vue'),
    ('Vue ref', 'camelCase', 'loading, tableData'),
    ('CSS类', 'kebab-case', 'proj-card'),
    ('API路径', '/api/{res}', '/api/purchase/orders'),
    ('JSON字段', 'snake_case', 'ship_status'),
    ('状态值', '中文', '已下单, 全部出货'),
]
for i, (a, b, c) in enumerate(data):
    r = t.rows[i+1].cells; r[0].text = a; r[1].text = b; r[2].text = c

doc.add_heading('八、反模式警示', level=1)
anti = [
    'API层写业务逻辑 -> 抽取到 services/',
    '多文件重复状态判断 -> 统一Service类',
    'el-table每次创建新对象 -> 预缓存 _enriched',
    '多表格selectedIds覆盖 -> 独立映射',
    'computed内_open状态 -> 独立Set管理',
    '前端硬编码page_size -> /all或统一1000',
    '新列不写迁移 -> ALTER TABLE加入init_db()',
]
for a in anti:
    doc.add_paragraph(a, style='List Bullet')

doc.add_heading('九、技术栈', level=1)
doc.add_paragraph('Python 3.10+ | FastAPI | SQLAlchemy | SQLite/MySQL\nVue 3 + Element Plus + Vite\nGitHub + Railway 自动部署')

doc.save(r'C:\Users\20210817\Desktop\MRP系统编码规范.docx')
print('OK')
